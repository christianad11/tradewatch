"""Build the two-page TradeWatch-LB capstone report from frozen local artifacts."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "TradeWatch-LB_Final_Report.docx"

FOREST = "07311F"
CEDAR = "145C3A"
LIME = "B8DB62"
INK = "163227"
MUTED = "526158"
PALE = "EFF4EB"
LINE = "CAD6CB"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid.gridCol_lst
    for col, width in zip(grid, widths, strict=True):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.1, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def add_text(doc, text, *, before=0, after=6, bold_prefix=None, italic=False):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=before, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_font(prefix, bold=True)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_font(remainder, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_font(run, italic=italic)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=12, after=5, line=1.0, keep=True)
    run = paragraph.add_run(text)
    set_font(run, size=13, color=CEDAR, bold=True)
    return paragraph


def add_kicker(doc, text):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=0, after=5, line=1.0, keep=True)
    run = paragraph.add_run(text.upper())
    set_font(run, size=8.5, color=CEDAR, bold=True)
    run.font.all_caps = True
    return paragraph


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [2340, 2340, 2340, 2340])
    metrics = [
        ("4,931", "MATCHED HS4 FLOWS"),
        ("$13.1B", "PAIRED TRADE VALUE"),
        ("62", "HIGH-PRIORITY LEADS"),
        ("100%", "TOP-1% SYNTHETIC RECALL"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics, strict=True):
        set_cell_shading(cell, PALE)
        first = cell.paragraphs[0]
        style_paragraph(first, before=0, after=2, line=1.0)
        run = first.add_run(value)
        set_font(run, size=18, color=FOREST, bold=True)
        second = cell.add_paragraph()
        style_paragraph(second, before=0, after=0, line=1.0)
        run = second.add_run(label)
        set_font(run, size=7.2, color=MUTED, bold=True)
    return table


def add_source_line(doc):
    p = doc.add_paragraph()
    style_paragraph(p, before=4, after=4, line=1.0)
    r = p.add_run("Source: Frozen TradeWatch-LB dashboard artifact; UN Comtrade public preview API; HS 2017 (H5), annual 2019–2021.")
    set_font(r, size=8.2, color=MUTED, italic=True)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=0, after=0, line=1.0)
    r = p.add_run("TRADEWATCH-LB  •  LEBNET AI FELLOWSHIP  •  RESEARCH PROTOTYPE")
    set_font(r, size=7.5, color=MUTED, bold=True)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, before=0, after=0, line=1.0)
    r = p.add_run("FINAL PROJECT REPORT  |  AUGUST 2026")
    set_font(r, size=7.5, color=MUTED, bold=True)


def set_document_geometry(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header(section)
    add_footer(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)


def build_report():
    doc = Document()
    set_document_geometry(doc)

    add_kicker(doc, "AI for Lebanon · Applied ML and public-interest data")
    title = doc.add_paragraph()
    style_paragraph(title, before=0, after=3, line=0.95, keep=True)
    run = title.add_run("TradeWatch-LB")
    set_font(run, size=25, color=FOREST, bold=True)
    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, before=0, after=11, line=1.05, keep=True)
    run = subtitle.add_run("Auditable machine-learning triage for unusual Lebanon mirror-trade flows")
    set_font(run, size=12, color=MUTED, italic=True)

    add_metric_strip(doc)
    add_source_line(doc)

    add_heading(doc, "Abstract")
    add_text(
        doc,
        "TradeWatch-LB is a research prototype that helps analysts prioritize unusual discrepancies between Lebanon’s reported imports and partner countries’ reported exports. It pairs official bilateral commodity records at the exact year, partner, and HS4 code; engineers interpretable signals; and combines a robust statistical baseline with two unsupervised anomaly detectors. In the frozen 2019–2021 MVP, the system scored 4,931 comparable flows, surfaced 62 high-priority investigation leads, and recovered all 80 severe synthetic perturbations in the top 1% of its test ranking. The system is deliberately not a fraud detector: every flag is an evidence file for human review, with caveats about legitimate sources of disagreement.",
    )

    add_heading(doc, "Problem and motivation")
    add_text(
        doc,
        "Trade statistics are valuable for oversight, but raw tables make it difficult to see which of thousands of bilateral product flows deserve attention. Mirror discrepancies can reflect real issues, yet they also arise from CIF/FOB valuation, shipping and reporting lags, re-exports, transit, partner attribution, commodity classification, estimation, and revisions. The project asks a narrower and safer question: where do Lebanon’s import records diverge unusually from partner-reported exports, after making the comparison transparent? The intended users are researchers, journalists, civic-data groups, and public-sector analysts who need a defensible triage view before deeper document-based investigation.",
    )

    add_heading(doc, "Data and method")
    add_text(
        doc,
        "The reproducible MVP caches responses from the official UN Comtrade public preview API. To avoid mixing product nomenclatures, it uses annual HS 2017 (H5) data for 2019–2021 and the top six Lebanese import partners in each year (eight unique partners). Lebanon import records are matched to partner export records only when year, partner, and HS4 code agree. The pipeline labels pairs Tier A when both value and quantity are comparable, Tier B when value alone is comparable, and Tier C when a mirror side is absent; Tier C is retained as a reporting-gap signal but never scored as a value discrepancy.",
    )
    add_text(
        doc,
        "For Tier A/B pairs, TradeWatch computes symmetric value gaps, signed log ratios, unit-value differences, within-series temporal shifts, and materiality. A median/MAD robust baseline, Isolation Forest, and Local Outlier Factor provide complementary anomaly evidence. Their rank-based outputs are combined into a transparent 0–100 Investigation Priority Score: mirror gap 30%, temporal change 20%, unit-value evidence 15%, model agreement 20%, and materiality 15%. The dashboard preserves both official records, component scores, comparability tier, and contextual caveats for every displayed lead.",
    )

    add_heading(doc, "Implementation and results")
    add_text(
        doc,
        "The implementation is a Python pipeline with cached raw JSON, deterministic feature engineering, a scored Parquet mirror-pair artifact, and a React evidence dashboard. The deployed interface includes a command center, filterable anomaly explorer, per-case evidence view, model-validation panel, methods/caveats page, and an accessible “How AI works” explainer. The final dataset contains 4,492 Tier A and 439 Tier B comparable records, covering 872 HS4 products. It also preserves 6,023 Tier C one-sided records separately rather than converting missingness into a discrepancy score.",
    )
    add_text(
        doc,
        "Because public trade statistics do not provide reliable labels for wrongdoing, the evaluation tests sensitivity rather than claimed fraud-detection accuracy. The pipeline selected 1,694 historically close mirror pairs, injected a one-sided 5× value perturbation into 80 of them, and assessed their percentile against the stable reference. The severe synthetic injections achieved 100% recall at both the top-1% and top-5% ranking cutoffs, with a mean percentile of 100.0. This demonstrates that the ranker responds strongly to a controlled, severe signal; it does not validate a causal or legal interpretation of real-world flags.",
    )

    add_heading(doc, "Discussion, limitations, and next steps")
    add_text(
        doc,
        "TradeWatch-LB succeeds as a reproducible investigation-prioritization MVP, not as an adjudication system. A high score is not a probability of fraud, smuggling, tax loss, corruption, or intent. The scoped preview API limited the first release to 2019–2021 and a partner sample, while bilateral data can be revised and lacks shipment-level documentation. The next iteration should use authenticated full API access, extend the time window without crossing HS revisions, audit partner and product coverage, incorporate documented customs context where legally appropriate, and test analyst usefulness through structured user studies. Any deployment should preserve the same human-review and non-accusatory guardrails.",
    )

    add_heading(doc, "Reflection on learning")
    add_text(
        doc,
        "This project showed why a smaller, auditable ML system can be more useful than an impressive-sounding black box. The rewarding part was turning a broad “AI for trade” idea into an interface that exposes its evidence and limits. The hardest parts were defining mirror-comparable records, avoiding incompatible HS revisions, and evaluating an unsupervised model without trustworthy fraud labels. The response was to make those constraints part of the product: cache the source data, label comparability, validate with a synthetic sensitivity test, and state what the model cannot conclude. That experience reinforced the need for transparent, human-in-the-loop AI systems that prioritize evidence over unsupported conclusions.",
    )

    add_heading(doc, "Reproducibility and sources")
    add_text(
        doc,
        "Code, cached acquisition metadata, model configuration, test suite, and dashboard artifact are included in the TradeWatch-LB repository. Primary data source: UN Comtrade public API (comtradeapi.un.org). Method reference: UN Comtrade, Methodology Guide for Comtrade Plus. This report and the repository README state the project scope, assumptions, and prohibited claims.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
